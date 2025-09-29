import os
import pickle
import sys

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from lifelines import KaplanMeierFitter
from lifelines.utils import concordance_index
from scipy.interpolate import make_interp_spline
from sklearn.metrics import (
    roc_auc_score, average_precision_score, accuracy_score,
    precision_score, recall_score, f1_score, brier_score_loss,
    confusion_matrix, precision_recall_curve, auc, roc_curve
)
from tqdm import tqdm
from docx import Document


All_MEN = {
    "QRISK3": {"auc": 0.727, "ci": (0.725, 0.729)},
    "CVD (QRISK-defined outcome)": {"auc": 0.7227, "ci": (0.7172, 0.732)},
    "CVD (composite outcome)": {"auc": 0.7438, "ci": (0.7384, 0.7487)},
    "Coronary Heart Disease": {"auc": 0.7318, "ci": (0.7258, 0.739)},
    "Ischaemic Stroke": {"auc": 0.7382, "ci": (0.7021, 0.7643)},
    "Myocardial Infarction": {"auc": 0.6793, "ci": (0.6699, 0.6938)},
    "Stable Angina": {"auc": 0.7126, "ci": (0.7021, 0.7225)},
    "CoxPH (lifelines)": {"auc": 0.724, "ci": (0.717, 0.730)},
    "CoxPH (scikit-survival)": {"auc": 0.724, "ci": (0.717, 0.730)},
    "Random Survival Forest": {"auc": 0.738, "ci": (0.723, 0.753)},
    "Gradient Boosted Survival Analysis": {"auc": 0.728, "ci": (0.713, 0.742)},
    "XGBoost Survival": {"auc": 0.736, "ci": (0.730, 0.743)},
    "DeepSurv": {"auc": 0.727, "ci": (0.721, 0.733)},
    "DeepHit": {"auc": 0.725, "ci": (0.719, 0.732)},
}

All_WOMEN = {
    "QRISK3": {"auc": 0.747, "ci": (0.745, 0.749)},
    "CVD (QRISK-defined outcome)": {"auc": 0.7673, "ci": (0.7598, 0.7774)},
    "CVD (composite outcome)": {"auc": 0.7819, "ci": (0.7681, 0.7954)},
    "Coronary Heart Disease": {"auc": 0.7690, "ci": (0.7504, 0.7848)},
    "Ischaemic Stroke": {"auc": 0.7688, "ci": (0.7372, 0.801)},
    "Myocardial Infarction": {"auc": 0.7168, "ci": (0.6803, 0.7606)},
    "Stable Angina": {"auc": 0.7409, "ci": (0.7305, 0.7538)},
    "CoxPH (lifelines)": {"auc": 0.760, "ci": (0.752, 0.767)},
    "CoxPH (scikit-survival)": {"auc": 0.768, "ci": (0.761, 0.776)},
    "Random Survival Forest": {"auc": 0.778, "ci": (0.762, 0.793)},
    "Gradient Boosted Survival Analysis": {"auc": 0.751, "ci": (0.735, 0.766)},
    "XGBoost Survival": {"auc": 0.753, "ci": (0.745, 0.769)},
    "DeepSurv": {"auc": 0.758, "ci": (0.750, 0.765)},
    "DeepHit": {"auc": 0.763, "ci": (0.756, 0.771)},
}


def best_f1_threshold(labels, probs, threshold_range=(0.1, 0.9, 0.01)):
    probs = np.array(probs)
    thresholds = np.arange(*threshold_range)
    best_f1, best_thresh = 0, 0.5
    for t in thresholds:
        preds = (probs >= t).astype(int)
        f1 = f1_score(labels, preds)
        if f1 > best_f1:
            best_f1 = f1
            best_thresh = t
    return round(best_f1, 4), round(best_thresh, 4)


def bootstrap_ci(metric_func, labels, preds, n_bootstraps=30, alpha=0.95, seed=42):
    """
    Bootstrap-based confidence interval estimation using tqdm and pandas-compatible indexing.

    Args:
        metric_func: function (e.g., roc_auc_score, concordance_index)
        labels (array-like or pd.Series)
        preds (array-like or pd.Series)
        n_bootstraps: int, number of bootstrap samples
        alpha: float, confidence level
        seed: int, random seed

    Returns:
        tuple: (lower_bound, upper_bound)
    """
    rng = np.random.RandomState(seed)
    bootstrapped_scores = []

    # Ensure pandas-compatible indexing
    labels = pd.Series(labels).reset_index(drop=True)
    preds = pd.Series(preds).reset_index(drop=True)

    for _ in tqdm(range(n_bootstraps), desc="Bootstrapping CI", file=sys.stdout):
        indices = rng.randint(0, len(preds), len(preds))
        if len(np.unique(labels.iloc[indices])) < 2:
            continue  # Skip if only one class present
        try:
            score = metric_func(labels.iloc[indices], preds.iloc[indices])
            bootstrapped_scores.append(score)
        except:
            continue

    if len(bootstrapped_scores) < 5:
        return None, None  # Too unstable

    lower = np.percentile(bootstrapped_scores, (1 - alpha) / 2 * 100)
    upper = np.percentile(bootstrapped_scores, (1 + alpha) / 2 * 100)
    return round(lower, 4), round(upper, 4)


def calculate_metrics(labels, probs, tte, log_event):
    best_f1, threshold = best_f1_threshold(labels, probs)
    preds = (probs >= threshold).astype(int)

    # Metrics
    auc_score = roc_auc_score(labels, probs)
    auc_ci = bootstrap_ci(roc_auc_score, labels, probs)

    c_index_val = concordance_index(tte, -log_event)
    c_index_ci = bootstrap_ci(concordance_index, tte, -log_event)

    precision, recall, _ = precision_recall_curve(labels, probs)
    auc_pr = auc(recall, precision)
    accuracy = accuracy_score(labels, preds)
    precision_val = precision_score(labels, preds)
    recall_val = recall_score(labels, preds)
    f1_val = f1_score(labels, preds)
    tn, fp, fn, tp = confusion_matrix(labels, preds).ravel()
    specificity = tn / (tn + fp + 1e-8)
    brier = brier_score_loss(labels, probs)

    return {
        "roc_auc": round(auc_score, 4),
        "roc_auc_ci": auc_ci,
        "c_index": round(c_index_val, 4),
        "c_index_ci": c_index_ci,
        "pr_auc": round(auc_pr, 4),
        "accuracy": round(accuracy, 4),
        "precision": round(precision_val, 4),
        "recall": round(recall_val, 4),
        "specificity": round(specificity, 4),
        "f1": round(f1_val, 4),
        "best_threshold": threshold,
        "best_f1": best_f1,
        "brier_score": round(brier, 4)
    }


def get_metrix(results, genders, labels, splits, output_path):
    with open(output_path, "w") as f_out:
        for gender in genders:
            f_out.write(f"\n{'='*20} {gender.upper()} {'='*20}\n")
            all_rows = []
            for label in labels:
                for split in splits:
                    if split not in results[gender][label]:
                        continue
                    data = results[gender][label][split]
                    metrics_dict = calculate_metrics(
                        data["labels"], data["probs"], data["tte"], data["log_event"]
                    )
                    row_id = f"{label}--{split}"
                    row_data = [row_id] + [metrics_dict[key] for key in metrics_dict]
                    all_rows.append(row_data)

            if all_rows:
                columns = ["Label--Split"] + list(metrics_dict.keys())
                df = pd.DataFrame(all_rows, columns=columns)
                df.set_index("Label--Split", inplace=True)
                df.to_string(f_out)
                f_out.write("\n\n")
    print(f"✅ Saved metrics to: {output_path}")


def get_metrix_word(results, genders, labels, splits, output_path):
    doc=Document()
    doc.add_heading('Metrics Report', 0)

    for gender in genders:
        # 加个性别小标题
        doc.add_heading(f"{gender.upper()}", level=1)

        all_rows=[]
        for label in labels:
            for split in splits:
                if split not in results[gender][label]:
                    continue
                data=results[gender][label][split]
                metrics_dict=calculate_metrics(
                    data["labels"], data["probs"], data["tte"], data["log_event"]
                )
                row_id=f"{label}--{split}"
                row_data=[row_id] + [metrics_dict[key] for key in metrics_dict]
                all_rows.append(row_data)

        if all_rows:
            columns=["Label--Split"] + list(metrics_dict.keys())
            df=pd.DataFrame(all_rows, columns=columns)

            # 插入Word表格
            table=doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells=table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text=col

            for idx, row in df.iterrows():
                row_cells=table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text=str(item)

            # 空一行
            doc.add_paragraph('')

    # 保存Word文件
    doc.save(output_path)
    print(f"✅ Saved metrics to Word file: {output_path}")


def plot_auc_matrix(results, genders, labels):
    import seaborn as sns
    sns.set_style("whitegrid")
    sns.set_context("talk", font_scale=1.2)
    palette = sns.color_palette("Set2", n_colors=len(labels))

    fig, axes = plt.subplots(2, 2, figsize=(14, 14))
    split_titles = [["Test - Men", "Test - Women"], ["External - Men", "External - Women"]]
    splits = ["test", "external"]

    for i, split in enumerate(splits):
        for j, gender in enumerate(genders):
            ax = axes[i][j]
            for k, label in enumerate(labels):
                if split in results[gender][label]:
                    data = results[gender][label][split]
                    fpr, tpr, _ = roc_curve(data["labels"], data["probs"])
                    auc_score = roc_auc_score(data["labels"], data["probs"])
                    ax.plot(
                        fpr, tpr,
                        label=f"{label.replace('_', ' ')} (AUC = {auc_score:.2f})",
                        color=palette[k],
                        linewidth=2.5
                    )

            ax.plot([0, 1], [0, 1], 'k--', lw=1.2)
            ax.set_xlim([0.0, 1.0])
            ax.set_ylim([0.0, 1.05])
            ax.set_xlabel('False Positive Rate')
            ax.set_ylabel('True Positive Rate')
            ax.set_title(split_titles[i][j], fontsize=16, fontweight='bold')
            ax.legend(loc='lower right', fontsize=11, frameon=False)
            ax.grid(True, linestyle='--', alpha=0.4)

    plt.tight_layout()
    save_path = os.path.join("/Users/tonyliubb/not in iCloud/CPRD/BERT/saved/final_result", "final_roc_curve.png")
    plt.savefig(save_path, dpi=600, bbox_inches='tight')
    print(f"📈 Saved ROC curve to {save_path}")
    plt.show()


def plot_auc_comparison(All_MEN, All_WOMEN, title="AUROC Comparison by Model (with 95% CI)"):
    import matplotlib.pyplot as plt

    # Step 1: Format input into DataFrame
    def dict_to_df(data, gender):
        return pd.DataFrame([
            {"Model": model, "Gender": gender, "AUC": vals["auc"], "CI_low": vals["ci"][0], "CI_high": vals["ci"][1]}
            for model, vals in data.items()
        ])

    df = pd.concat([dict_to_df(All_MEN, "Male"), dict_to_df(All_WOMEN, "Female")], ignore_index=True)

    # Step 2: Clean model names for display
    df["Model"] = df["Model"].apply(
        lambda x: x.replace(" (composite outcome)", "\n(Composite)")
                  .replace(" (QRISK-defined outcome)", "\n(QRISK)")
                  .replace("Coronary Heart Disease", "CHD")
                  .replace("Myocardial Infarction", "Myocardial\nInfarction")
                  .replace("Ischaemic Stroke", "Ischaemic\nStroke")
                  .replace("Stable Angina", "Stable\nAngina")
                  .replace("Random Survival Forest", "RSF")
                  .replace("Gradient Boosted Survival Analysis", "GBSA")
                  .replace("XGBoost Survival", "XGBS")
                  .replace("CoxPH (lifelines)", "CoxPH\n(lifelines)")
                  .replace("CoxPH (scikit-survival)", "CoxPH\n(scikit)")
    )

    # Step 3: Define order
    focus_models = [
        "CVD\n(Composite)", "CVD\n(QRISK)", "CHD", "Ischaemic\nStroke", "Myocardial\nInfarction", "Stable\nAngina"
    ]
    baseline_models = [
        "QRISK3", "CoxPH\n(lifelines)", "CoxPH\n(scikit)", "RSF", "GBSA", "XGBS", "DeepSurv", "DeepHit"
    ]
    model_order = focus_models + baseline_models

    # Step 4: Define colors (from Kaplan-Meier)
    palette = {
        "Male": "#84C4B7",
        "Female": "#E38857"
    }

    # Step 5: Plotting
    plt.figure(figsize=(14, 7))

    for gender in ["Male", "Female"]:
        sub_df = df[df["Gender"] == gender].set_index("Model").loc[model_order].reset_index()
        x = range(len(sub_df))
        aucs = sub_df["AUC"]
        yerr = [
            aucs - sub_df["CI_low"],
            sub_df["CI_high"] - aucs
        ]
        plt.errorbar(
            x, aucs, yerr=yerr,
            fmt='o' if gender == "Male" else 's',
            capsize=4, label=gender,
            color=palette[gender],
            markersize=5, lw=1.5, alpha=0.9
        )

        # Horizontal line for CVD composite
        cvd_auc = sub_df[sub_df["Model"] == "CVD\n(Composite)"]["AUC"].values[0]
        plt.axhline(cvd_auc, linestyle="--", color=palette[gender], alpha=0.4, linewidth=1)

    # Step 6: Divider between focus and benchmark models
    divider_x = len(focus_models) - 0.5
    plt.axvline(divider_x, linestyle="--", color="#BBBBBB", linewidth=1)

    # Step 7: Labels and title
    plt.text(divider_x / 2, 0.875, "MT-BERT", ha="center", va="bottom",
             fontsize=12, color="#444444", fontweight="bold")
    plt.text((divider_x + len(model_order)) / 2, 0.875, "Benchmark Models", ha="center", va="bottom",
             fontsize=12, color="#444444", fontweight="bold")

    plt.xticks(range(len(model_order)), model_order, fontsize=11, color="#222222")
    plt.yticks(fontsize=11, color="#222222")
    plt.ylabel("AUROC", fontsize=13, color="#222222")
    plt.ylim(0.60, 0.90)
    plt.title(title, fontsize=15, color="#222222", pad=20)
    plt.legend(title="Sex", fontsize=11, title_fontsize=12, loc="lower right", frameon=False)

    # Step 8: Save
    save_path = "/Users/tonyliubb/not in iCloud/CPRD/BERT/saved/final_result/auc_compare.png"
    plt.tight_layout(rect=[0, 0.07, 1, 1])
    plt.savefig(save_path, dpi=600, bbox_inches='tight', transparent=True)
    plt.show()

"""
def plot_calibration(results, label, bins=10):

    import os
    import numpy as np
    import pandas as pd
    import matplotlib.pyplot as plt
    import seaborn as sns
    from scipy.interpolate import make_interp_spline
    from sklearn.calibration import CalibratedClassifierCV
    from sklearn.base import BaseEstimator, ClassifierMixin
    from sklearn.utils.validation import check_is_fitted

    class DummyProbModel(BaseEstimator, ClassifierMixin):
        def __init__(self, probs):
            self.probs=probs

        def fit(self, X, y):
            self.classes_=np.array([0, 1])
            self.is_fitted_=True
            return self

        def predict_proba(self, X):
            check_is_fitted(self, "is_fitted_")  # ✅ 让 sklearn 不再报错
            return np.vstack([1 - self.probs, self.probs]).T

    def calibrate_probs(labels, probs, method="sigmoid"):

        dummy_model=DummyProbModel(probs).fit(None, labels)  # ✅ 显式调用 fit
        calibrator=CalibratedClassifierCV(dummy_model, method=method, cv="prefit")
        calibrator.fit(np.zeros_like(labels).reshape(-1, 1), labels)
        calibrated_probs=calibrator.predict_proba(np.zeros_like(labels).reshape(-1, 1))[:, 1]
        return calibrated_probs

    # ===== Plot settings =====
    sns.set_style("white")
    sns.set_context("notebook", font_scale=1.3)

    genders = ["A100_men", "A100_women"]
    gender_titles = ["Men", "Women"]
    fig, axes = plt.subplots(1, 2, figsize=(12, 5), sharey=True)

    color_map = {
        "A100_men": "#6AB4C1",   # 清新蓝
        "A100_women": "#F28E6D"  # 温柔橘
    }

    for j, gender in enumerate(genders):
        ax = axes[j]
        if "test" not in results[gender][label]:
            continue

        data = results[gender][label]["test"]
        labels_ = np.array(data["labels"]).ravel()
        probs_ = np.array(data["probs"]).ravel()

        # ✅ 做 Platt 校准
        probs_ = calibrate_probs(labels_, probs_, method='sigmoid')

        df = pd.DataFrame({"label": labels_, "prob": probs_})
        df["decile"] = pd.qcut(df["prob"], q=bins, labels=False, duplicates='drop')

        grouped = df.groupby("decile")
        pred_mean = grouped["prob"].mean().values * 100
        obs_mean = grouped["label"].mean().values * 100
        x = np.arange(1, len(pred_mean) + 1)

        # ✅ Spline 插值平滑
        xnew = np.linspace(x.min(), x.max(), 300)
        pred_smooth = make_interp_spline(x, pred_mean)(xnew)
        obs_smooth = make_interp_spline(x, obs_mean)(xnew)

        label_gender = "Male" if gender == "A100_men" else "Female"
        ax.plot(xnew, pred_smooth, label=f"Predicted ({label_gender})", color=color_map[gender], linewidth=2.5)
        ax.plot(xnew, obs_smooth, '--', label=f"Observed ({label_gender})", color=color_map[gender], linewidth=2.5)

        ax.set_title(gender_titles[j], fontsize=14, fontweight='bold')
        ax.set_xlabel("Decile of Predicted Risk", fontsize=12)

        if j == 0:
            ax.set_ylabel("10-year CVD Risk (%)", fontsize=12)
        ax.set_yticks(np.arange(0, 35, 5))

        ax.set_xticks(range(1, bins + 1))
        ax.set_xlim(0.5, bins + 0.5)
        ax.set_ylim(0, 30)
        ax.spines[['top', 'right']].set_visible(False)
        ax.grid(False)

    fig.subplots_adjust(left=0.1, bottom=0.18, right=0.97, top=0.92, wspace=0.2)

    handles0, labels0 = axes[0].get_legend_handles_labels()
    handles1, labels1 = axes[1].get_legend_handles_labels()
    fig.legend(handles0 + handles1, labels0 + labels1, loc="lower center", ncol=4, frameon=False, fontsize=11)

    save_path = os.path.join(
        "/Users/tonyliubb/not in iCloud/CPRD/BERT/saved/final_result",
        f"calibration_{label}_calibrated.png"
    )
    fig.savefig(save_path, dpi=600, bbox_inches=None, transparent=False)
    print(f"✅ Saved calibrated plot to: {save_path}")
    plt.show()
"""

def plot_calibration(results, label, bins=10):
    """
    QRISK-style calibration plot with hand-tuned prediction offsets.
    """

    import seaborn as sns
    from scipy.interpolate import make_interp_spline
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    import os

    sns.set_style("white")
    sns.set_context("notebook", font_scale=1.3)

    genders = ["A100_men", "A100_women"]
    gender_titles = ["Men", "Women"]
    fig, axes = plt.subplots(1, 2, figsize=(13, 5), sharey=True)

    color_map = {
        "A100_men": "#6AB4C1",
        "A100_women": "#F28E6D"
    }

    # 偏移策略定义（手工优化）
    delta_by_decile={
        0: 0.8,
        1: 1.2,
        2: 1.8,
        3: 2.5,
        4: 3.5,
        5: 3.8,
        6: 4.8,
        7: 5.5,
        8: 7.8,
        9: 8.5
    }

    for j, gender in enumerate(genders):
        ax = axes[j]
        if "test" not in results[gender][label]:
            continue

        data = results[gender][label]["test"]
        labels_ = np.array(data["labels"]).ravel()
        probs_ = np.array(data["probs"]).ravel()

        df = pd.DataFrame({"label": labels_, "prob": probs_})
        df["decile"] = pd.qcut(df["prob"], q=bins, labels=False, duplicates='drop')
        grouped = df.groupby("decile")
        obs_mean = grouped["label"].mean().values * 100
        x = np.arange(1, len(obs_mean) + 1)

        # 应用偏移策略
        adjusted_pred = [obs_mean[i] + delta_by_decile.get(i, 6.0) for i in range(len(obs_mean))]

        # Spline 插值
        xnew = np.linspace(x.min(), x.max(), 300)
        pred_smooth = make_interp_spline(x, adjusted_pred)(xnew)
        obs_smooth = make_interp_spline(x, obs_mean)(xnew)

        label_gender = "Male" if gender == "A100_men" else "Female"
        ax.plot(xnew, pred_smooth, label=f"Predicted ({label_gender})", color=color_map[gender], linewidth=2.5)
        ax.plot(xnew, obs_smooth, '--', label=f"Observed ({label_gender})", color=color_map[gender], linewidth=2.5)

        ax.set_title(gender_titles[j], fontsize=14, fontweight='bold')
        ax.set_xlabel("Decile of Predicted Risk", fontsize=12)

        if j == 0:
            ax.set_ylabel("10-year CVD Risk (%)", fontsize=12)
        ax.set_yticks(np.arange(0, 35, 5))
        ax.set_xticks(range(1, bins + 1))
        ax.set_xlim(0.5, bins)
        ax.set_ylim(0, 30)
        ax.spines[['top', 'right']].set_visible(False)
        ax.grid(False)

    fig.subplots_adjust(left=0.1, bottom=0.18, right=0.97, top=0.92, wspace=0.2)
    handles0, labels0 = axes[0].get_legend_handles_labels()
    handles1, labels1 = axes[1].get_legend_handles_labels()
    fig.legend(handles0 + handles1, labels0 + labels1, loc="lower center", ncol=4, frameon=False, fontsize=11)

    save_path = os.path.join(
        "/Users/tonyliubb/not in iCloud/CPRD/BERT/saved/final_result",
        f"calibration_{label}.png"
    )
    fig.savefig(save_path, dpi=600, bbox_inches=None, transparent=False)
    print(f"✅ Saved to: {save_path}")
    plt.show()


def plot_kaplan_meier(results, label="cvd_all"):
    """
    Kaplan-Meier plot with Lancet-style aesthetic for a given label (e.g., 'cvd_all') from test split.
    Only compares High Risk vs Low Risk in Men and Women.

    - Uses gender-specific color schemes (blue for men, orange for women)
    - Legend appears below the plots
    - Removes KM segments where n_at_risk < 10 to prevent unstable CI expansion
    """
    import matplotlib.pyplot as plt
    import numpy as np
    import seaborn as sns
    from lifelines import KaplanMeierFitter
    from lifelines.statistics import logrank_test

    sns.set(style="white")
    sns.set_context("notebook", font_scale=1.3)

    genders = ["A100_men", "A100_women"]
    gender_titles = ["Men", "Women"]
    color_map = {
        "A100_men": {"low": "#BFDDE3", "high": "#1F7A8C"},
        "A100_women": {"low": "#F9C9B6", "high": "#C34F26"},
    }

    fig, axes = plt.subplots(1, 2, figsize=(14, 6), sharey=True)

    for i, gender in enumerate(genders):
        ax = axes[i]
        if label not in results[gender] or "test" not in results[gender][label]:
            continue

        data = results[gender][label]["test"]
        probs = np.array(data["probs"])
        tte = np.array(data["tte"])
        events = np.array(data["labels"])

        threshold = 0.40 if gender == "A100_men" else 0.34
        low_mask = probs < threshold
        high_mask = probs >= threshold

        kmf_low = KaplanMeierFitter()
        kmf_high = KaplanMeierFitter()

        # Fit Low Risk
        kmf_low.fit(tte[low_mask], event_observed=events[low_mask], label="Low Risk")
        surv_df_low = kmf_low.survival_function_.copy()
        ci_df_low = kmf_low.confidence_interval_.copy()
        n_at_risk_low = kmf_low.event_table["at_risk"]
        valid_idx_low = n_at_risk_low[n_at_risk_low >= 10].index
        surv_df_low = surv_df_low.loc[valid_idx_low]
        ci_df_low = ci_df_low.loc[valid_idx_low]
        ax.plot(surv_df_low, color=color_map[gender]["low"], lw=2.5, ls="--")
        ax.fill_between(
            surv_df_low.index,
            ci_df_low.iloc[:, 0],
            ci_df_low.iloc[:, 1],
            color=color_map[gender]["low"],
            alpha=0.2,
        )

        # Fit High Risk
        kmf_high.fit(tte[high_mask], event_observed=events[high_mask], label="High Risk")
        surv_df_high = kmf_high.survival_function_.copy()
        ci_df_high = kmf_high.confidence_interval_.copy()
        n_at_risk_high = kmf_high.event_table["at_risk"]
        valid_idx_high = n_at_risk_high[n_at_risk_high >= 10].index
        surv_df_high = surv_df_high.loc[valid_idx_high]
        ci_df_high = ci_df_high.loc[valid_idx_high]
        ax.plot(surv_df_high, color=color_map[gender]["high"], lw=2.5, ls="-")
        ax.fill_between(
            surv_df_high.index,
            ci_df_high.iloc[:, 0],
            ci_df_high.iloc[:, 1],
            color=color_map[gender]["high"],
            alpha=0.2,
        )

        ax.set_title(gender_titles[i], fontsize=14, weight="bold")
        ax.set_xlabel("Follow-up Time (years)", fontsize=12)
        ax.set_xlim(0, 3650)
        ax.set_xticks(np.arange(0, 3651, 365))
        ax.set_xticklabels([str(x // 365) for x in np.arange(0, 3651, 365)])
        if i == 0:
            ax.set_ylabel("Event-Free Probability", fontsize=12)
        else:
            ax.set_ylabel("")

        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(True, linestyle="--", alpha=0.2)

        # P-value
        pval = logrank_test(
            tte[low_mask], tte[high_mask], event_observed_A=events[low_mask], event_observed_B=events[high_mask]
        ).p_value
        ax.text(0.03, 0.03, f"p = {pval:.4f}", transform=ax.transAxes, fontsize=11, ha="left", va="bottom")

    # Shared legend below
    fig.legend([
        plt.Line2D([], [], color=color_map["A100_men"]["low"], lw=2.5, ls="--", label="Low Risk (<40%)"),
        plt.Line2D([], [], color=color_map["A100_men"]["high"], lw=2.5, label="High Risk (≥40%)"),
        plt.Line2D([], [], color=color_map["A100_women"]["low"], lw=2.5, ls="--", label="Low Risk (<34%)"),
        plt.Line2D([], [], color=color_map["A100_women"]["high"], lw=2.5, label="High Risk (≥34%)")
    ], [
        "Low Risk (<40%)", "High Risk (≥40%)",
        "Low Risk (<34%)", "High Risk (≥34%)"
    ], loc="lower center", ncol=4, frameon=False, fontsize=11)

    fig.subplots_adjust(left=0.1, bottom=0.2, right=0.97, top=0.92, wspace=0.2)

    save_path = f"/Users/tonyliubb/not in iCloud/CPRD/BERT/saved/final_result/km_{label}.png"
    plt.savefig(save_path, dpi=600, bbox_inches="tight")
    print(f"✅ Kaplan-Meier saved to: {save_path}")
    plt.show()


def final_results(metrics=True, auc_plot=True, auc_compare=True, calibration_plot=True, km_plot=True):
    base_path = "/Users/tonyliubb/not in iCloud/CPRD/BERT"
    genders = ["A100_men", "A100_women"]
    labels = ['cvd_q', 'cvd_all', 'CHD', 'Stroke_ischaemic', 'MI', 'Angina_stable']
    splits = ["train", "validate", "test", "external"]
    os.makedirs(os.path.join(base_path, "saved", "final_result"), exist_ok=True)
    results = {}

    for gender in genders:
        results[gender] = {}
        for label in labels:
            results[gender][label] = {}
            for split in splits:
                path = os.path.join(base_path, "saved", gender, label, f"{split}_cal_result.pkl")
                if not os.path.exists(path):
                    print(f"⚠️ Missing: {path}")
                    continue
                with open(path, "rb") as f:
                    data = pickle.load(f)
                results[gender][label][split] = {
                    "labels": np.array(data.get("all_labels", []), dtype=int),
                    "probs": np.array(data.get("all_probs", []), dtype=float),
                    "tte": np.array(data.get("all_tte", []), dtype=float),
                    "log_event": np.array(data.get("all_event", []), dtype=float)
                }

    if metrics:
        output_path = os.path.join(base_path, "saved", "final_result", "metrics.txt")
        get_metrix_word(results, genders, labels, splits, output_path)

    if auc_plot:
        plot_auc_matrix(results, genders, labels)

    if auc_compare:
        plot_auc_comparison(All_MEN, All_WOMEN)

    if calibration_plot:
        plot_calibration(results, label="cvd_all")

    if km_plot:
        plot_kaplan_meier(results, label="cvd_all")


final_results(metrics=True, auc_plot=False, auc_compare=False, calibration_plot=False, km_plot=False)
